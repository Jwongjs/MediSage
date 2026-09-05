from typing import Dict, Any
from datetime import datetime
import logging

# PDF/Word imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)


def _explanation_text(explanation) -> str | None:
    """Explanation may be a dataclass instance (raw graph state) or a dict
    (already-serialized state) -- same dual-type tolerance as `canonical`
    below."""
    if explanation is None:
        return None
    text = explanation["text"] if isinstance(explanation, dict) else explanation.text
    return text or None


class MedicalReportNode:
    """Renders a finished diagnosis state as a downloadable document.

    MediSage stores nothing, so this node has no database dependency and no
    persistence methods: it only turns in-memory graph state into bytes that
    the API streams straight back to the caller.
    """

#==================================================
# Medical Report Export Function
#==================================================

    # Export functionality (only called from API endpoint)
    async def generate_export_file(self, state: dict, format: str, include_details: bool = True) -> bytes:
        """Generate PDF or Word export file (separate from main workflow)"""

        if format == 'pdf':
            return await self._generate_pdf_export(state, include_details)
        elif format == 'word':
            return await self._generate_word_export(state, include_details)
        else:
            raise ValueError(f"Unsupported format: {format}")

    async def _generate_pdf_export(self, state: dict, include_details: bool) -> bytes:
        """Generate PDF from the evidence state (ranking/matrix/canonical/judgements)."""
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=18)
            styles = getSampleStyleSheet()
            story = []

            story.append(Paragraph("MediSage — Differential Summary", styles['Title']))
            story.append(Spacer(1, 12))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}", styles['Normal']))
            story.append(Paragraph(f"Session ID: {state.get('session_id', 'Unknown')}", styles['Normal']))
            story.append(Spacer(1, 12))

            story.append(Paragraph("Reported Information", styles['Heading2']))
            story.append(Paragraph(state.get("patient_text", "(none)"), styles['Normal']))
            story.append(Spacer(1, 12))

            canonical = {c["key"] if isinstance(c, dict) else c.key:
                         (c["label"] if isinstance(c, dict) else c.label)
                         for c in state.get("canonical", [])}
            judgements = state.get("judgements", {})
            matrix = state.get("matrix", {})
            explanations = state.get("explanations", {})

            for group in state.get("ranking", []):
                tied = len(group) > 1
                for diagnosis in group:
                    heading_text = f"{diagnosis}  (tied)" if tied else diagnosis
                    story.append(Paragraph(heading_text, styles['Heading2']))

                    definition = _explanation_text(explanations.get(diagnosis))
                    if definition:
                        story.append(Paragraph(definition, styles['Normal']))
                        story.append(Spacer(1, 6))

                    for status, heading in (
                        ("supported", "Supported by"),
                        ("contradicted", "Contradicted by"),
                        ("not_mentioned", "Not yet established"),
                    ):
                        rows = [
                            (canonical.get(k, k), importance, judgements.get(k, {}).get("evidence"))
                            for k, importance in matrix.get(diagnosis, {}).items()
                            if judgements.get(k, {}).get("status", "not_mentioned") == status
                        ]
                        if not rows:
                            continue
                        story.append(Paragraph(heading, styles['Heading3']))
                        for label, importance, evidence in rows:
                            story.append(Paragraph(f"[{importance}] {label}", styles['Normal']))
                            if evidence and include_details:
                                story.append(Paragraph(f'Patient said: "{evidence}"', styles['Normal']))
                    story.append(Spacer(1, 12))

            not_evaluated = state.get("not_evaluated", [])
            if not_evaluated:
                story.append(Paragraph("Considered but not assessed", styles['Heading2']))
                story.append(Paragraph(
                    "These candidates could not be evaluated against the available evidence and are not ranked:",
                    styles['Normal']))
                for name in not_evaluated:
                    story.append(Paragraph(f"- {name}", styles['Normal']))
                story.append(Spacer(1, 12))

            summary = state.get("summary") or {}
            if summary:
                story.append(Paragraph("Clinical Summary", styles['Heading2']))
                if summary.get("severity"):
                    story.append(Paragraph(f"Severity: {summary['severity'].title()}", styles['Normal']))
                if summary.get("specialist_recommendation"):
                    story.append(Paragraph(
                        f"Recommended specialist: {summary['specialist_recommendation'].replace('_', ' ').title()}",
                        styles['Normal']))
                if summary.get("user_explanation"):
                    story.append(Paragraph(summary["user_explanation"], styles['Normal']))
                    if summary.get("explanation_source"):
                        source_line = f"Source: {summary['explanation_source']}"
                        if summary.get("explanation_url"):
                            source_line += f" ({summary['explanation_url']})"
                        story.append(Paragraph(source_line, styles['Normal']))
                story.append(Spacer(1, 12))

            # Always include disclaimer
            story.append(Paragraph("Disclaimer", styles['Heading3']))
            story.append(Paragraph(
                "This is not a diagnosis. Share it with a healthcare professional.",
                styles['Normal']))

            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"PDF generation failed: {e}")
            # Fallback to text
            return self._generate_text_export(state, include_details).encode('utf-8')

    async def _generate_word_export(self, state: Dict[str, Any], include_details: bool) -> bytes:
        """Generate a Word document from the evidence state (ranking/matrix/canonical/judgements)."""
        try:
            doc = Document()

            title = doc.add_heading('MediSage — Differential Summary', 0)
            title.alignment = 1  # Center alignment

            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
            doc.add_paragraph(f"Session ID: {state.get('session_id', 'Unknown')}")

            doc.add_heading('Reported Information', level=1)
            doc.add_paragraph(state.get("patient_text", "(none)"))

            canonical = {c["key"] if isinstance(c, dict) else c.key:
                         (c["label"] if isinstance(c, dict) else c.label)
                         for c in state.get("canonical", [])}
            judgements = state.get("judgements", {})
            matrix = state.get("matrix", {})
            explanations = state.get("explanations", {})

            for group in state.get("ranking", []):
                tied = len(group) > 1
                for diagnosis in group:
                    heading_text = f"{diagnosis}  (tied)" if tied else diagnosis
                    doc.add_heading(heading_text, level=1)

                    definition = _explanation_text(explanations.get(diagnosis))
                    if definition:
                        doc.add_paragraph(definition)

                    for status, heading in (
                        ("supported", "Supported by"),
                        ("contradicted", "Contradicted by"),
                        ("not_mentioned", "Not yet established"),
                    ):
                        rows = [
                            (canonical.get(k, k), importance, judgements.get(k, {}).get("evidence"))
                            for k, importance in matrix.get(diagnosis, {}).items()
                            if judgements.get(k, {}).get("status", "not_mentioned") == status
                        ]
                        if not rows:
                            continue
                        doc.add_heading(heading, level=2)
                        for label, importance, evidence in rows:
                            doc.add_paragraph(f"[{importance}] {label}", style='List Bullet')
                            if evidence and include_details:
                                doc.add_paragraph(f'Patient said: "{evidence}"')

            not_evaluated = state.get("not_evaluated", [])
            if not_evaluated:
                doc.add_heading('Considered but not assessed', level=1)
                doc.add_paragraph(
                    "These candidates could not be evaluated against the available evidence and are not ranked:")
                for name in not_evaluated:
                    doc.add_paragraph(name, style='List Bullet')

            summary = state.get("summary") or {}
            if summary:
                doc.add_heading('Clinical Summary', level=1)
                if summary.get("severity"):
                    doc.add_paragraph(f"Severity: {summary['severity'].title()}")
                if summary.get("specialist_recommendation"):
                    doc.add_paragraph(
                        f"Recommended specialist: {summary['specialist_recommendation'].replace('_', ' ').title()}")
                if summary.get("user_explanation"):
                    doc.add_paragraph(summary["user_explanation"])
                    if summary.get("explanation_source"):
                        source_line = f"Source: {summary['explanation_source']}"
                        if summary.get("explanation_url"):
                            source_line += f" ({summary['explanation_url']})"
                        doc.add_paragraph(source_line)

            # Always include disclaimer
            doc.add_heading('Disclaimer', level=2)
            doc.add_paragraph("This is not a diagnosis. Share it with a healthcare professional.")

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Word generation failed: {e}")
            # Fallback to text
            return self._generate_text_export(state, include_details).encode('utf-8')

    def _generate_text_export(self, state: dict, include_details: bool = True) -> str:
        lines = ["MEDISAGE — DIFFERENTIAL SUMMARY", "=" * 60, ""]
        lines.append("REPORTED INFORMATION")
        lines.append(state.get("patient_text", "(none)"))
        lines.append("")

        canonical = {c["key"] if isinstance(c, dict) else c.key:
                     (c["label"] if isinstance(c, dict) else c.label)
                     for c in state.get("canonical", [])}
        judgements = state.get("judgements", {})
        matrix = state.get("matrix", {})
        explanations = state.get("explanations", {})

        for group in state.get("ranking", []):
            tied = len(group) > 1
            for diagnosis in group:
                lines.append(f"{diagnosis}{'  (tied)' if tied else ''}")
                lines.append("-" * 60)
                definition = _explanation_text(explanations.get(diagnosis))
                if definition:
                    lines.append(definition)
                    lines.append("")
                for status, heading in (
                    ("supported", "Supported by"),
                    ("contradicted", "Contradicted by"),
                    ("not_mentioned", "Not yet established"),
                ):
                    rows = [
                        (canonical.get(k, k), importance, judgements.get(k, {}).get("evidence"))
                        for k, importance in matrix.get(diagnosis, {}).items()
                        if judgements.get(k, {}).get("status", "not_mentioned") == status
                    ]
                    if not rows:
                        continue
                    lines.append(f"  {heading}:")
                    for label, importance, evidence in rows:
                        lines.append(f"    - [{importance}] {label}")
                        if evidence and include_details:
                            lines.append(f"        patient said: \"{evidence}\"")
                lines.append("")

        not_evaluated = state.get("not_evaluated", [])
        if not_evaluated:
            lines.append("CONSIDERED BUT NOT ASSESSED (not ranked)")
            lines.append("-" * 60)
            lines.append("These candidates could not be evaluated against the available evidence:")
            for name in not_evaluated:
                lines.append(f"  - {name}")
            lines.append("")

        lines.append("This is not a diagnosis. Share it with a healthcare professional.")
        return "\n".join(lines)
